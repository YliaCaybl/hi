import requests
from bs4 import BeautifulSoup
from docx import Document

# URL сайта
url = 'https://eco.lenobl.ru/ru/deiatelnost/prognoz-neblagopriyatnykh-meteorologicheskikh-uslovij-leningradskoj-ob/informaciya-ot-fgbu-severo-zapadnoe-ugms/'

# Путь к файлу Word
output_path = r'C:\Users\yioly\pars_628910\output.docx'

# Загрузка страницы
response = requests.get(url)
soup = BeautifulSoup(response.content, 'html.parser')

# Поиск таблицы
table = soup.find('table', border="1", cellpadding="1", cellspacing="1", style="width:500px")

# Поиск последнего тега <tr> в таблице
if table:
    last_tr = table.find_all('tr')[-1]
    last_tr_text = ' '.join(td.get_text(strip=True) for td in last_tr.find_all('td'))
else:
    last_tr_text = 'Таблица не найдена'

# Открытие существующего документа Word
doc = Document(output_path)

# Добавление ссылки на сайт в начало документа
doc.add_paragraph(url)

# Добавление текста из последнего тега <tr>
doc.add_paragraph(last_tr_text)

# Сохранение документа
doc.save(output_path)

print(f'Текст успешно добавлен в {output_path}')
