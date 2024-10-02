import requests
from bs4 import BeautifulSoup
from docx import Document
import os

# URL сайта
url = 'http://pogoda-sv.ru/pollcenter/nmu/'

# Путь к документу
doc_path = r'C:\Users\yioly\pars_628910\output.docx'


def extract_text_from_first_p(soup):
    # Ищем первый <p> с атрибутом style="text-align:justify"
    p_tag = soup.find('p', style='text-align:justify')
    if p_tag:
        # Получаем текст
        return p_tag.get_text(strip=True)
    return ""


def remove_text_after_phrase(text, phrase="Предупреждение переданона предприятия"):
    # Находим позицию фразы
    phrase_index = text.find(phrase)
    if phrase_index != -1:
        # Возвращаем текст до фразы
        return text[:phrase_index].strip()
    return text


try:
    # Получение содержимого сайта
    headers = {
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/58.0.3029.110 Safari/537.3'}
    response = requests.get(url, headers=headers)
    response.raise_for_status()  # Проверка на ошибки HTTP
    soup = BeautifulSoup(response.content, 'html.parser')

    # Извлечение текста из первого <p> с нужным стилем
    full_text = extract_text_from_first_p(soup)

    # Удаление части текста после фразы "Предупреждение передано на предприятия"
    cleaned_text = remove_text_after_phrase(full_text)

    # Проверка, существует ли файл
    if os.path.exists(doc_path):
        # Открытие существующего документа
        doc = Document(doc_path)
    else:
        # Создание нового документа
        doc = Document()

    # Добавление заголовка с URL (всегда добавляется)
    doc.add_heading('Сайт: ' + url, level=1)

    # Добавление текста после удаления части
    if cleaned_text:
        doc.add_paragraph(f"\n{cleaned_text}", style='Normal')

    # Сохранение документа
    doc.save(doc_path)

    print("Текст успешно добавлен в документ.")

except requests.RequestException as e:
    print(f"Ошибка при запросе данных: {e}")
except Exception as e:
    print(f"Произошла ошибка: {e}")


kyjgtliktugvlbiuyliuvy
