import requests
from bs4 import BeautifulSoup
from docx import Document
import os

# URL сайта
url = 'http://www.ecomos.ru/kadr22/predupOwuzWnmuMos.asp'

# Путь к документу
doc_path = r'C:\Users\yioly\pars_628910\output.docx'


def extract_text_from_td_tags(soup):
    # Ищем все <td> с классом fon22
    td_tags = soup.find_all('td', class_='fon22')
    text_parts = []

    for td in td_tags:
        # Удаляем все <p> теги и их содержимое
        for p in td.find_all('p'):
            p.decompose()
        # Удаляем <table> теги и их содержимое
        for table in td.find_all('table'):
            table.decompose()
        # Удаляем <div> с id="disswfl" и его содержимое
        diss_div = td.find('div', id='disswfl')
        if diss_div:
            diss_div.decompose()
        # Собираем текст из <td>, оставшийся после удаления <p>, <table> и <div>
        text = td.get_text(strip=True)
        if text:
            text_parts.append(text)

    return "\n".join(text_parts)


try:
    # Получение содержимого сайта
    headers = {
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/58.0.3029.110 Safari/537.3'}
    response = requests.get(url, headers=headers)
    response.raise_for_status()  # Проверка на ошибки HTTP
    soup = BeautifulSoup(response.content, 'html.parser')

    # Поиск текста из <td class="fon22">
    text_to_add = extract_text_from_td_tags(soup)

    # Проверка, существует ли файл
    if os.path.exists(doc_path):
        # Открытие существующего документа
        doc = Document(doc_path)
    else:
        # Создание нового документа
        doc = Document()

    # Добавление заголовка с URL (всегда добавляется)
    doc.add_heading('Сайт: ' + url, level=1)

    # Добавление текста из <td class="fon22">
    doc.add_paragraph(text_to_add, style='Normal')

    # Сохранение документа
    doc.save(doc_path)

    print("Текст успешно добавлен в документ.")

except requests.RequestException as e:
    print(f"Ошибка при запросе данных: {e}")
except Exception as e:
    print(f"Произошла ошибка: {e}")
